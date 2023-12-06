from docx import Document
from docx.shared import Pt
from transformers import GPT2Tokenizer, GPT2LMHeadModel

# for loading the model and tokenizer 
tokenizer = GPT2Tokenizer.from_pretrained('gpt2', pad_token='[PAD]', padding_side='left')
model = GPT2LMHeadModel.from_pretrained('gpt2')

def extend_content(doc_path: str, tokenizer, model, min_lines: int = 10, max_title_length: int = 50, max_lines_per_page: int = 50) -> Document:
    doc = Document(doc_path)
    processed_paragraphs = set() 

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() and paragraph.text not in processed_paragraphs:  
            lines = paragraph.text.split('\n')

            if len(lines) < min_lines or len(paragraph.text) < 50 or "specific_keyword" not in paragraph.text:
                try:
                    input_ids = tokenizer.encode(paragraph.text, return_tensors='pt')
                    attention_mask = input_ids.ne(tokenizer.pad_token_id).float()

                    output = model.generate(
                        input_ids,
                        attention_mask=attention_mask,
                        max_length=80,
                        temperature=2.0,  # increase for more output
                        top_k=50,
                        top_p=0.9
                    )

                    extra_content = tokenizer.decode(output[0], skip_special_tokens=True)
                    extra_lines = extra_content.split('\n')

                    #for estimation if the page is filled
                    if len(lines) + len(extra_lines) > max_lines_per_page:
                        extra_content = '\n'.join(extra_lines[:max_lines_per_page - len(lines)])

                    run = paragraph.add_run(extra_content)
                    run.font.size = Pt(12)

                    processed_paragraphs.add(paragraph.text)  # adding paragraph 

                except RuntimeError as e:
                    print(f"A RuntimeError occurred: {e}")
                except ValueError as e:
                    print(f"A ValueError occurred: {e}")
                except Exception as e:
                    print(f"An unexpected error occurred: {e}")

    return doc

doc_path = "output.docx" #input document
extended_doc = extend_content(doc_path, tokenizer, model)
extended_doc.save("extended_document.docx") #output document
