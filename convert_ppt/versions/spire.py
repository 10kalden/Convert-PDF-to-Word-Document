import os
from PIL import Image as PilImage
import pytesseract
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from spire.presentation.common import *
from spire.presentation import *

try:
    if not os.path.exists('Output'):
        os.makedirs('Output')

    presentation = Presentation()
    presentation.LoadFromFile("Unit1.pptx")
    doc = Document()

    for i, slide in enumerate(presentation.Slides):
        fileName = "Output/Image_" + str(i) + ".png"
        image = slide.SaveAsImage()
        image.Save(fileName)
        
        img = PilImage.open(fileName)
        text = pytesseract.image_to_string(img)
        doc.add_picture(fileName, width=Inches(5))
        doc.add_paragraph(text)
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        image.Dispose()

    doc.save("Output/docs.docx")
    presentation.Dispose()

except Exception as e:
    print(f"An error occurred: {e}")
